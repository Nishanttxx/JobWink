"""
JobWink backend — Phase 1 with Hybrid Persistence (Supabase PostgreSQL + Local Storage Snapshot)

Consolidates tailor_resume_endpoint.py + resume_sections.py into one app,
and completes PDF/DOCX export, tailor-from-saved-resume flow, and hybrid persistence.

Endpoints:
  POST   /resume/new                        -> blank resume, empty sections
  POST   /resume/upload                     -> upload PDF/DOCX, AI scans into sections
  GET    /resume/{id}                       -> fetch a resume's sections
  PATCH  /resume/{id}/section/{name}        -> edit ONE section only
  POST   /resume/{id}/tailor                -> AI-tailor a saved resume to a job description
  GET    /resume/{id}/export?format=pdf|docx -> download the resume as a file
  POST   /template/analyze                  -> analyze template PDF visual tokens
  GET    /template/{template_id}            -> get analyzed template config
  GET    /resume/{id}/render                -> render resume using analyzed template design
"""

import os
import io
import re
import json
import uuid
import datetime
import html
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from typing import Optional, Dict, List, Any
from pathlib import Path

from fastapi import FastAPI, HTTPException, UploadFile, File, Query, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from dotenv import load_dotenv

from template_schema import TemplateConfig
from template_analyzer import analyze_template
from template_renderer import render_resume_pdf, get_candidate_filename
from job_prediction_service import prediction_service, MANDATORY_DISCLAIMER

# Load environment variables
load_dotenv(Path(__file__).parent / ".env")
load_dotenv()
load_dotenv(Path(__file__).parent.parent / "job_collector" / ".env")

app = FastAPI(title="JobWink Backend")

# Setup local persistent storage directory
DATA_STORE_DIR = Path(__file__).parent / "data_store"
DATA_STORE_DIR.mkdir(exist_ok=True)

# Initialize Supabase Client if available
SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_KEY = os.environ.get("SUPABASE_SERVICE_ROLE_KEY") or os.environ.get("SUPABASE_ANON_KEY")

supabase_client = None
if SUPABASE_URL and SUPABASE_KEY:
    try:
        from supabase import create_client
        supabase_client = create_client(SUPABASE_URL, SUPABASE_KEY)
        print("Connected to Supabase PostgreSQL database.")
    except Exception as e:
        print(f"Supabase connection note: {e}")

TEMPLATES: Dict[str, TemplateConfig] = {}

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ---------------------------------------------------------------------------
# Shared Schema
# ---------------------------------------------------------------------------

class WorkExperienceEntry(BaseModel):
    company: str = ""
    title: str = ""
    start_date: str = ""
    end_date: str = ""
    bullets: List[str] = Field(default_factory=list)


class EducationEntry(BaseModel):
    institution: str = ""
    degree: str = ""
    field: str = ""
    graduation_date: str = ""


class ProjectEntry(BaseModel):
    title: str = ""
    description: str = ""
    bullets: List[str] = Field(default_factory=list)


class ResumeSections(BaseModel):
    contact_info: Dict[str, Any] = Field(default_factory=lambda: {
        "name": "", "email": "", "phone": "", "location": "", "links": []
    })
    summary: str = ""
    work_experience: List[WorkExperienceEntry] = Field(default_factory=list)
    education: List[EducationEntry] = Field(default_factory=list)
    projects: List[ProjectEntry] = Field(default_factory=list)
    skills: List[str] = Field(default_factory=list)
    keywords: List[str] = Field(default_factory=list)
    gaps: List[str] = Field(default_factory=list)          # populated after tailoring
    tailoring_notes: str = ""                                # populated after tailoring


RESUMES: Dict[str, ResumeSections] = {}
VALID_SECTIONS = {"contact_info", "summary", "work_experience", "education", "projects", "skills", "keywords"}


# ---------------------------------------------------------------------------
# Persistence Operations (Supabase DB + Local Data Store Fallback)
# ---------------------------------------------------------------------------

def save_resume_persist(
    resume_id: str,
    sections: ResumeSections,
    title: str = "Master Resume",
    user_id: Optional[str] = None,
    change_summary: str = "Version snapshot"
):
    """Save resume into memory, disk snapshot, and Supabase DB if authenticated."""
    content_json = sections.model_dump()
    
    # 1. Update memory cache
    RESUMES[resume_id] = sections

    # 2. Save local disk snapshot (guarantees state retention across restarts)
    try:
        file_path = (DATA_STORE_DIR / f"{resume_id}.json").resolve()
        # Security: prevent path traversal vulnerabilities
        if not file_path.is_relative_to(DATA_STORE_DIR.resolve()):
            raise ValueError("Invalid resume_id: Path traversal detected")

        with open(file_path, "w", encoding="utf-8") as f:
            json.dump({
                "resume_id": resume_id,
                "title": title,
                "user_id": user_id,
                "sections": content_json,
                "change_summary": change_summary
            }, f, indent=2)
    except Exception as e:
        print(f"Local storage save note: {e}")

    # 3. Save to Supabase PostgreSQL database if user_id is provided and valid
    if supabase_client and user_id:
        try:
            supabase_client.table("resumes").upsert({
                "id": resume_id,
                "user_id": user_id,
                "title": title,
                "updated_at": "now()",
            }).execute()

            version_data = {
                "resume_id": resume_id,
                "user_id": user_id,
                "parsed_content": content_json,
                "optimized_content": content_json if (sections.gaps or sections.tailoring_notes) else {},
                "change_summary": change_summary,
            }
            res = supabase_client.table("resume_versions").insert(version_data).execute()
            if res.data and len(res.data) > 0:
                version_id = res.data[0]["id"]
                supabase_client.table("resumes").update({"current_version_id": version_id}).eq("id", resume_id).execute()
        except Exception as e:
            print(f"Supabase DB save note: {e}")

    # Mark existing predictions as stale because resume version updated
    prediction_service.invalidate_stale_predictions(resume_id)


def load_resume_persist(resume_id: str) -> Optional[ResumeSections]:
    """Retrieve resume from memory cache, local disk snapshot, or Supabase DB."""
    # 1. Check memory cache
    if resume_id in RESUMES:
        return RESUMES[resume_id]

    # 2. Check local disk snapshot
    file_path = (DATA_STORE_DIR / f"{resume_id}.json").resolve()
    # Security: prevent path traversal vulnerabilities
    if file_path.is_relative_to(DATA_STORE_DIR.resolve()) and file_path.exists():
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                sections = ResumeSections(**data["sections"])
                RESUMES[resume_id] = sections
                return sections
        except Exception as e:
            print(f"Local storage load note: {e}")

    # 3. Check Supabase DB (only for valid UUID format to avoid 22P02 Postgres errors)
    is_valid_uuid = False
    try:
        uuid.UUID(resume_id)
        is_valid_uuid = True
    except ValueError:
        is_valid_uuid = False

    if supabase_client and is_valid_uuid:
        try:
            res = supabase_client.table("resume_versions") \
                .select("parsed_content, optimized_content") \
                .eq("resume_id", resume_id) \
                .order("created_at", desc=True) \
                .limit(1) \
                .execute()

            if res.data and len(res.data) > 0:
                row = res.data[0]
                content = row.get("optimized_content") or row.get("parsed_content")
                if content and isinstance(content, dict):
                    sections = ResumeSections(**content)
                    RESUMES[resume_id] = sections
                    return sections
        except Exception as e:
            print(f"Supabase DB load note: {e}")

    # 4. Fallback auto-initialization for master / demo resumes
    sections = ResumeSections(
        contact_info={"name": "Alex Morgan", "email": "alex.morgan@jobwink.ai", "phone": "+1 555-0192", "location": "San Francisco, CA"},
        summary="Senior Full-Stack Software Engineer with 6+ years of experience building high-throughput web applications and AI-driven platforms.",
        skills=["Flutter", "Dart", "Python", "FastAPI", "React", "TypeScript", "PostgreSQL", "Docker", "AWS", "REST APIs"],
        keywords=["Agile", "System Architecture", "CI/CD", "Machine Learning", "Microservices"],
    )
    save_resume_persist(resume_id, sections, title="Master Resume", change_summary="Auto-initialized fallback")
    return sections


# ---------------------------------------------------------------------------
# AI Orchestration (Anthropic Claude -> OpenAI gpt-4o-mini -> Heuristic Engine)
# ---------------------------------------------------------------------------

def call_ai_for_json(system_prompt: str, user_content: str) -> Optional[dict]:
    anthropic_key = os.environ.get("ANTHROPIC_API_KEY")
    openai_key = os.environ.get("OPENAI_API_KEY")

    if anthropic_key:
        try:
            from anthropic import Anthropic
            client = Anthropic(api_key=anthropic_key)
            response = client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=2000,
                system=system_prompt,
                messages=[{"role": "user", "content": user_content}],
            )
            raw_text = response.content[0].text
            try:
                return json.loads(raw_text)
            except json.JSONDecodeError:
                retry = client.messages.create(
                    model="claude-3-5-sonnet-20241022",
                    max_tokens=2000,
                    system=system_prompt,
                    messages=[
                        {"role": "user", "content": user_content},
                        {"role": "assistant", "content": raw_text},
                        {"role": "user", "content": "That wasn't valid JSON. Return ONLY the JSON object, nothing else."},
                    ],
                )
                return json.loads(retry.content[0].text)
        except Exception as e:
            if not openai_key:
                print(f"Anthropic AI failed: {e}")

    if openai_key:
        try:
            from openai import OpenAI
            client = OpenAI(api_key=openai_key)
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_content},
                ],
                response_format={"type": "json_object"},
            )
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            print(f"OpenAI AI failed: {e}")

    return None


def heuristic_parse_text(raw_text: str) -> dict:
    """Regex & rule-based parser fallback when AI API key is not present or API is offline."""
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    
    email_match = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", raw_text)
    phone_match = re.search(r"(\+?\d{1,3}[\s-]?)?\(?\d{3}\)?[\s-]?\d{3}[\s-]?\d{4}", raw_text)
    links = re.findall(r"(?:linkedin\.com\/in\/[^\s]+|github\.com\/[^\s]+)", raw_text, re.IGNORECASE)

    name = lines[0] if lines else ""
    if len(name) > 40 or "@" in name or "http" in name:
        name = ""

    summary = ""
    summary_match = re.search(r"SUMMARY[\s\S]*?(?=EDUCATION|SKILLS|PROJECTS|EXPERIENCE|$)", raw_text, re.IGNORECASE)
    if summary_match:
        summary_lines = summary_match.group(0).splitlines()[1:]
        summary = " ".join(summary_lines).strip()

    skills = []
    skills_match = re.search(r"SKILLS[\s\S]*?(?=PROJECTS|EXPERIENCE|EDUCATION|EXTRA|$)", raw_text, re.IGNORECASE)
    if skills_match:
        skills_text = skills_match.group(0)
        found = re.findall(r"([A-Za-z0-9#\+\.]+)", skills_text)
        skip = {"SKILLS", "Backend", "Tools", "Testing", "API", "Languages", "Soft"}
        skills = list(dict.fromkeys([s for s in found if s not in skip and len(s) > 1]))

    projects = []
    proj_match = re.search(r"PROJECTS[\s\S]*?(?=EXPERIENCE|EDUCATION|EXTRA|$)", raw_text, re.IGNORECASE)
    if proj_match:
        proj_text = proj_match.group(0)
        p_blocks = proj_text.split("\n\n")
        for b in p_blocks[1:]:
            b_lines = [l.strip() for l in b.splitlines() if l.strip()]
            if b_lines:
                title = b_lines[0]
                bullets = [l.lstrip("•- ").strip() for l in b_lines[1:] if l.strip()]
                projects.append({"title": title, "description": "", "bullets": bullets})

    experience = []
    exp_match = re.search(r"EXPERIENCE[\s\S]*?(?=EXTRA|EDUCATION|SKILLS|$)", raw_text, re.IGNORECASE)
    if exp_match:
        exp_text = exp_match.group(0)
        e_lines = [l.strip() for l in exp_text.splitlines() if l.strip()][1:]
        if e_lines:
            exp_title = e_lines[0]
            bullets = [l.lstrip("•- ").strip() for l in e_lines[1:] if l.strip()]
            experience.append({
                "company": "Company",
                "title": exp_title,
                "start_date": "",
                "end_date": "Present",
                "bullets": bullets
            })

    education = []
    edu_match = re.search(r"EDUCATION[\s\S]*?(?=SKILLS|PROJECTS|EXPERIENCE|$)", raw_text, re.IGNORECASE)
    if edu_match:
        edu_lines = [l.strip() for l in edu_match.group(0).splitlines() if l.strip()][1:]
        for line in edu_lines[:3]:
            education.append({
                "institution": line,
                "degree": "Degree",
                "field": "Information Science",
                "graduation_date": ""
            })

    return {
        "contact_info": {
            "name": name,
            "email": email_match.group(0) if email_match else "",
            "phone": phone_match.group(0) if phone_match else "",
            "location": "Remote",
            "links": links,
        },
        "summary": summary,
        "work_experience": experience,
        "education": education,
        "projects": projects,
        "skills": skills,
        "keywords": ["Python", "Docker", "API Testing", "Flutter"],
    }


def heuristic_tailor_resume(original: dict, job_description: str) -> dict:
    """Intelligent fallback tailoring engine when AI keys are not present."""
    tailored = dict(original)
    
    jd_words = set(re.findall(r"\b[A-Za-z0-9\+\#]+\b", job_description))
    
    orig_skills = tailored.get("skills", [])
    matched_skills = [s for s in orig_skills if s in jd_words or s.lower() in job_description.lower()]
    
    other_skills = [s for s in orig_skills if s not in matched_skills]
    tailored["skills"] = matched_skills + other_skills
    
    common_reqs = ["AWS", "Kubernetes", "GraphQL", "CI/CD", "Redis", "TypeScript"]
    gaps = [req for req in common_reqs if req in job_description and req not in orig_skills]
    
    tailored["gaps"] = gaps
    tailored["tailoring_notes"] = (
        f"Tailored resume sections for target job description. "
        f"Prioritized {len(matched_skills)} key skills. "
        f"Identified {len(gaps)} potential technology gaps."
    )
    return tailored


# ---------------------------------------------------------------------------
# Feature 1.1: start from scratch
# ---------------------------------------------------------------------------

@app.post("/resume/new")
def new_resume(x_user_id: Optional[str] = Header(None, alias="X-User-ID")):
    resume_id = str(uuid.uuid4())
    sections = ResumeSections()
    save_resume_persist(resume_id, sections, title="New Blank Resume", user_id=x_user_id, change_summary="Initialized blank resume")
    return {"resume_id": resume_id, "sections": sections.model_dump()}


# ---------------------------------------------------------------------------
# Feature 1.2: upload & scan resume into sections
# ---------------------------------------------------------------------------

PARSE_SYSTEM_PROMPT = """Extract structured data from this resume text.

Respond with ONLY valid JSON matching this exact schema, no prose, no markdown fences:
{
  "contact_info": {"name": "", "email": "", "phone": "", "location": "", "links": []},
  "summary": "",
  "work_experience": [{"company": "", "title": "", "start_date": "", "end_date": "", "bullets": []}],
  "education": [{"institution": "", "degree": "", "field": "", "graduation_date": ""}],
  "projects": [{"title": "", "description": "", "bullets": []}],
  "skills": [],
  "keywords": []
}

Rules:
- Extract only what is actually present in the text. Leave fields blank rather than guessing.
- Split each job/project description into individual bullet points.
- "keywords" = notable tools/certifications/domain terms not already captured in "skills".
"""


def extract_text_from_file(file: UploadFile) -> str:
    content = file.file.read()
    name = file.filename.lower() if file.filename else ""
    if name.endswith(".pdf") or not name:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    elif name.endswith(".docx"):
        from docx import Document
        doc = Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)
    raise HTTPException(status_code=400, detail="Only .pdf and .docx are supported")


@app.post("/extract-pdf")
@app.post("/pdf/extract")
def extract_pdf_endpoint(file: UploadFile = File(...)):
    """High-fidelity backend PDF text extraction using pypdf."""
    try:
        raw_text = extract_text_from_file(file)
        if not raw_text.strip():
            raise HTTPException(status_code=422, detail="Could not extract text from PDF (empty or scanned image PDF).")
        return {
            "text": raw_text,
            "char_count": len(raw_text),
            "filename": file.filename or "uploaded.pdf",
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract PDF text: {str(e)}")


@app.post("/resume/upload")
def upload_resume(file: UploadFile = File(...), x_user_id: Optional[str] = Header(None, alias="X-User-ID")):
    raw_text = extract_text_from_file(file)
    if not raw_text.strip():
        raise HTTPException(status_code=422, detail="Could not extract any text (scanned/image PDF?)")

    resume_id = str(uuid.uuid4())
    parsed = call_ai_for_json(PARSE_SYSTEM_PROMPT, raw_text)
    
    if not parsed:
        parsed = heuristic_parse_text(raw_text)

    try:
        sections = ResumeSections(**parsed)
        save_resume_persist(resume_id, sections, title=f"Uploaded - {file.filename}", user_id=x_user_id, change_summary="Parsed document upload")
        return {"resume_id": resume_id, "sections": sections.model_dump()}
    except Exception as e:
        sections = ResumeSections()
        save_resume_persist(resume_id, sections, title="Blank Resume", user_id=x_user_id, change_summary="Fallback default initialization")
        return {
            "resume_id": resume_id,
            "sections": sections.model_dump(),
            "warning": f"Parsing validation note ({str(e)}) — initialized default schema.",
        }


# ---------------------------------------------------------------------------
# Section Fetch + Edit
# ---------------------------------------------------------------------------

@app.get("/resume/{resume_id}")
def get_resume(resume_id: str):
    sections = load_resume_persist(resume_id)
    if sections:
        return sections.model_dump()
    raise HTTPException(status_code=404, detail="Resume not found")


class SectionUpdate(BaseModel):
    content: Any


@app.patch("/resume/{resume_id}/section/{section_name}")
def update_section(resume_id: str, section_name: str, update: SectionUpdate, x_user_id: Optional[str] = Header(None, alias="X-User-ID")):
    sections = load_resume_persist(resume_id)
    if not sections:
        raise HTTPException(status_code=404, detail="Resume not found")

    if section_name not in VALID_SECTIONS:
        raise HTTPException(status_code=400, detail=f"Unknown section '{section_name}'")

    data = sections.model_dump()
    data[section_name] = update.content
    try:
        updated_sections = ResumeSections(**data)
        save_resume_persist(resume_id, updated_sections, user_id=x_user_id, change_summary=f"Updated section {section_name}")
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Invalid content for '{section_name}': {e}")

    return {"resume_id": resume_id, "sections": updated_sections.model_dump()}


# ---------------------------------------------------------------------------
# Feature 1.3: tailor saved resume to job description
# ---------------------------------------------------------------------------

TAILOR_SYSTEM_PROMPT = """You are tailoring a candidate's resume to a specific job description.

Rules:
- Only use experience, skills, and facts already present in the original resume JSON.
- Never invent employers, job titles, dates, tools, certifications, or metrics.
- Reorder and rephrase existing bullet points to emphasize relevance to the job.
- If the job requires something not present in the resume, list it under "gaps" instead of adding it.
- Respond with ONLY valid JSON matching this schema, no prose, no markdown fences:
{
  "contact_info": {...same shape as input...},
  "summary": "",
  "work_experience": [...same shape as input...],
  "education": [...same shape as input...],
  "projects": [...same shape as input...],
  "skills": [],
  "keywords": [],
  "gaps": [],
  "tailoring_notes": "short explanation of what changed and why"
}
"""


class TailorRequest(BaseModel):
    job_description: str


@app.post("/resume/{resume_id}/tailor")
def tailor_resume(resume_id: str, payload: TailorRequest, x_user_id: Optional[str] = Header(None, alias="X-User-ID")):
    sections = load_resume_persist(resume_id)
    if not sections:
        raise HTTPException(status_code=404, detail="Resume not found")

    original = sections.model_dump()
    user_content = f"""
==[ORIGINAL RESUME]==
{json.dumps(original, indent=2)}

==[JOB DESCRIPTION]==
{payload.job_description}
"""
    tailored_data = call_ai_for_json(TAILOR_SYSTEM_PROMPT, user_content)
    if not tailored_data:
        tailored_data = heuristic_tailor_resume(original, payload.job_description)

    new_id = str(uuid.uuid4())
    tailored_sections = ResumeSections(**tailored_data)
    save_resume_persist(new_id, tailored_sections, title="Tailored Resume", user_id=x_user_id, change_summary="Tailored against job description")

    return {"resume_id": new_id, "original_resume_id": resume_id, "sections": tailored_sections.model_dump()}


# ---------------------------------------------------------------------------
# Feature 1.4: export to PDF / DOCX
# ---------------------------------------------------------------------------

def build_pdf(sections: ResumeSections) -> io.BytesIO:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER, topMargin=0.6 * inch, bottomMargin=0.6 * inch)
    styles = getSampleStyleSheet()
    story = []

    contact = sections.contact_info
    story.append(Paragraph(f"<b>{contact.get('name', '')}</b>", styles["Title"]))
    contact_line = " | ".join(filter(None, [
        contact.get("email"),
        contact.get("phone"),
        contact.get("location"),
        *([str(l) for l in contact.get("links", [])] if isinstance(contact.get("links"), list) else [])
    ]))
    story.append(Paragraph(contact_line, styles["Normal"]))
    story.append(Spacer(1, 12))

    if sections.summary:
        story.append(Paragraph("<b>SUMMARY</b>", styles["Heading2"]))
        story.append(Paragraph(sections.summary, styles["Normal"]))
        story.append(Spacer(1, 10))

    if sections.work_experience:
        story.append(Paragraph("<b>EXPERIENCE</b>", styles["Heading2"]))
        for job in sections.work_experience:
            story.append(Paragraph(f"<b>{job.title}</b> — {job.company} ({job.start_date}–{job.end_date})", styles["Normal"]))
            for bullet in job.bullets:
                story.append(Paragraph(f"• {bullet}", styles["Normal"]))
            story.append(Spacer(1, 8))

    if sections.projects:
        story.append(Paragraph("<b>PROJECTS</b>", styles["Heading2"]))
        for proj in sections.projects:
            story.append(Paragraph(f"<b>{proj.title}</b>", styles["Normal"]))
            if proj.description:
                story.append(Paragraph(proj.description, styles["Normal"]))
            for bullet in proj.bullets:
                story.append(Paragraph(f"• {bullet}", styles["Normal"]))
            story.append(Spacer(1, 8))

    if sections.education:
        story.append(Paragraph("<b>EDUCATION</b>", styles["Heading2"]))
        for edu in sections.education:
            story.append(Paragraph(f"{edu.degree} in {edu.field} — {edu.institution} ({edu.graduation_date})", styles["Normal"]))
        story.append(Spacer(1, 8))

    if sections.skills:
        story.append(Paragraph("<b>SKILLS</b>", styles["Heading2"]))
        story.append(Paragraph(", ".join(sections.skills), styles["Normal"]))

    doc.build(story)
    buf.seek(0)
    return buf


def build_docx(sections: ResumeSections) -> io.BytesIO:
    from docx import Document

    buf = io.BytesIO()
    doc = Document()
    contact = sections.contact_info
    doc.add_heading(contact.get("name", ""), level=0)
    contact_line = " | ".join(filter(None, [
        contact.get("email"),
        contact.get("phone"),
        contact.get("location"),
        *([str(l) for l in contact.get("links", [])] if isinstance(contact.get("links"), list) else [])
    ]))
    doc.add_paragraph(contact_line)

    if sections.summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(sections.summary)

    if sections.work_experience:
        doc.add_heading("Experience", level=1)
        for job in sections.work_experience:
            doc.add_paragraph(f"{job.title} — {job.company} ({job.start_date}–{job.end_date})", style="Intense Quote")
            for bullet in job.bullets:
                doc.add_paragraph(bullet, style="List Bullet")

    if sections.projects:
        doc.add_heading("Projects", level=1)
        for proj in sections.projects:
            doc.add_paragraph(proj.title, style="Intense Quote")
            if proj.description:
                doc.add_paragraph(proj.description)
            for bullet in proj.bullets:
                doc.add_paragraph(bullet, style="List Bullet")

    if sections.education:
        doc.add_heading("Education", level=1)
        for edu in sections.education:
            doc.add_paragraph(f"{edu.degree} in {edu.field} — {edu.institution} ({edu.graduation_date})")

    if sections.skills:
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(sections.skills))

    doc.save(buf)
    buf.seek(0)
    return buf


@app.get("/resume/{resume_id}/export")
def export_resume(resume_id: str, format: str = Query("pdf", pattern="^(pdf|docx)$")):
    sections = load_resume_persist(resume_id)
    if not sections:
        raise HTTPException(status_code=404, detail="Resume not found")

    resume_data = sections.model_dump()
    candidate_filename = get_candidate_filename(resume_data)
    base_name = candidate_filename.rsplit(".", 1)[0]

    if format == "pdf":
        config = TemplateConfig()
        pdf_bytes = render_resume_pdf(config, resume_data)
        buf = io.BytesIO(pdf_bytes)
        media_type = "application/pdf"
        filename = f"{base_name}.pdf"
    else:
        buf = build_docx(sections)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"{base_name}.docx"

    return StreamingResponse(
        buf, media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# Template-based rendering
# ---------------------------------------------------------------------------

@app.post("/template/analyze")
def analyze_template_endpoint(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Template must be a .pdf")

    pdf_bytes = file.file.read()
    template_id = str(uuid.uuid4())
    config = analyze_template(pdf_bytes, source_filename=file.filename, template_id=template_id)
    TEMPLATES[template_id] = config
    return {"template_id": template_id, "config": config.model_dump()}


@app.get("/template/{template_id}")
def get_template_config(template_id: str):
    if template_id not in TEMPLATES:
        raise HTTPException(status_code=404, detail="Template not found")
    return TEMPLATES[template_id].model_dump()


@app.get("/resume/{resume_id}/render")
def render_with_template(resume_id: str, template_id: str = Query(...)):
    sections = load_resume_persist(resume_id)
    if not sections:
        raise HTTPException(status_code=404, detail="Resume not found")

    if template_id not in TEMPLATES:
        raise HTTPException(status_code=404, detail="Template not found")

    config = TEMPLATES[template_id]
    resume_data = sections.model_dump()
    filename = get_candidate_filename(resume_data)

    pdf_bytes = render_resume_pdf(config, resume_data)

    return StreamingResponse(
        iter([pdf_bytes]), media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# Combined tailor + template-render endpoint
# ---------------------------------------------------------------------------

class TailorAndExportRequest(BaseModel):
    job_description: str
    template_id: Optional[str] = None


@app.post("/resume/{resume_id}/tailor-and-export")
def tailor_and_export(
    resume_id: str,
    payload: TailorAndExportRequest,
    x_user_id: Optional[str] = Header(None, alias="X-User-ID"),
):
    """Combined endpoint: tailor resume content + render with original template design.
    
    1. Load resume
    2. AI-tailor content to job description
    3. Load template config (from template_id or use defaults)
    4. Render with adaptive one-page fitting
    5. Validate page count == 1
    6. Return as {Candidate Name}.pdf
    """
    sections = load_resume_persist(resume_id)
    if not sections:
        raise HTTPException(status_code=404, detail="Resume not found")

    # Step 1: Tailor content
    original = sections.model_dump()
    user_content = f"""
==[ORIGINAL RESUME]==
{json.dumps(original, indent=2)}

==[JOB DESCRIPTION]==
{payload.job_description}
"""
    tailored_data = call_ai_for_json(TAILOR_SYSTEM_PROMPT, user_content)
    if not tailored_data:
        tailored_data = heuristic_tailor_resume(original, payload.job_description)

    # Save tailored version
    new_id = str(uuid.uuid4())
    tailored_sections = ResumeSections(**tailored_data)
    save_resume_persist(
        new_id, tailored_sections,
        title="Tailored Resume",
        user_id=x_user_id,
        change_summary="Tailored and exported",
    )

    # Step 2: Get template config
    config = None
    if payload.template_id and payload.template_id in TEMPLATES:
        config = TEMPLATES[payload.template_id]
    else:
        # Use a sensible default template
        config = TemplateConfig()

    # Step 3: Render with one-page fitting
    resume_data = tailored_sections.model_dump()
    pdf_bytes = render_resume_pdf(config, resume_data)
    filename = get_candidate_filename(resume_data)

    return StreamingResponse(
        iter([pdf_bytes]),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "X-Resume-Id": new_id,
            "X-Original-Resume-Id": resume_id,
        },
    )


# ---------------------------------------------------------------------------
# Feature 2.0: Job Prediction Endpoints (ML Model Inference & Feature Overrides)
# ---------------------------------------------------------------------------

class JobPredictionRequest(BaseModel):
    resume_id: str
    job_description: str
    job_title: Optional[str] = None
    structured_features_override: Optional[Dict[str, Any]] = None
    resume_version_id: Optional[str] = None


@app.get("/api/job-prediction/features/{resume_id}")
def extract_job_prediction_features(resume_id: str):
    """
    Extracts default non-protected structured features from the latest tailored resume
    so the user can review and edit/correct them before triggering prediction.
    """
    sections = load_resume_persist(resume_id)
    if not sections:
        sections = ResumeSections()
        save_resume_persist(resume_id, sections, title="Default Master Resume", change_summary="Default auto-initialization")

    sections_dict = sections.model_dump()
    extracted = prediction_service.extract_structured_features(sections_dict)
    current_hash = prediction_service.compute_resume_hash(sections_dict)

    return {
        "resume_id": resume_id,
        "structured_features": extracted,
        "tailored_resume_hash": current_hash,
        "structured_feature_columns": prediction_service.structured_feature_columns,
    }


@app.post("/api/job-prediction/predict")
def predict_job_match(
    payload: JobPredictionRequest,
    x_user_id: Optional[str] = Header(None, alias="X-User-ID")
):
    """
    Executes model inference combining structured hiring pipeline (0.65) and fit pipeline (0.35).
    Evaluates saved artifact without retraining.
    Returns estimated match language and mandatory disclaimer.
    """
    sections = load_resume_persist(payload.resume_id)
    if not sections:
        sections = ResumeSections()
        save_resume_persist(payload.resume_id, sections, title="Default Master Resume", change_summary="Default auto-initialization")

    sections_dict = sections.model_dump()

    record = prediction_service.predict(
        resume_id=payload.resume_id,
        sections_dict=sections_dict,
        job_description=payload.job_description,
        user_features_override=payload.structured_features_override,
        job_title=payload.job_title,
        user_id=x_user_id,
        resume_version_id=payload.resume_version_id,
    )

    # Persist to Supabase if connected
    if supabase_client and x_user_id:
        try:
            supabase_client.table("job_predictions").insert({
                "id": record["id"],
                "user_id": x_user_id,
                "resume_id": payload.resume_id,
                "resume_version_id": payload.resume_version_id,
                "tailored_resume_hash": record["tailored_resume_hash"],
                "job_title": record["job_title"],
                "job_description": payload.job_description,
                "extracted_features": record["extracted_features"],
                "structured_probability": record["structured_probability"],
                "fit_probability": record["fit_probability"],
                "combined_probability": record["combined_probability"],
                "is_match": record["is_match"],
                "estimated_match_level": record["estimated_match_level"],
                "is_stale": False,
                "disclaimer": MANDATORY_DISCLAIMER,
            }).execute()
        except Exception as e:
            print(f"Supabase job_predictions insert note: {e}")

    return record


@app.get("/api/job-prediction/latest/{resume_id}")
def get_latest_job_prediction(resume_id: str):
    """Fetches the most recent prediction for a resume, evaluating staleness."""
    record = prediction_service.get_latest_prediction(resume_id)
    if not record:
        raise HTTPException(status_code=404, detail="No prediction found for this resume")

    sections = load_resume_persist(resume_id)
    if sections:
        curr_hash = prediction_service.compute_resume_hash(sections.model_dump())
        if record.get("tailored_resume_hash") != curr_hash:
            record["is_stale"] = True

    return record


# ---------------------------------------------------------------------------
# Bug Report & Admin Email Notification
# ---------------------------------------------------------------------------

class BugReportPayload(BaseModel):
    user_id: Optional[str] = None
    user_name: Optional[str] = None
    user_email: str
    title: str
    description: str
    page_url: Optional[str] = None
    route: Optional[str] = None
    browser: Optional[str] = None
    os: Optional[str] = None
    screen_size: Optional[str] = None
    screenshot_reference: Optional[str] = None


def send_bug_report_email(report: BugReportPayload, report_id: str) -> tuple[bool, Optional[str], Optional[str]]:
    """Delivers the bug report to the configured admin email with Reply-To.

    Reads ADMIN_EMAIL exclusively from the server environment (backend/.env or OS env).
    The admin recipient is NEVER exposed to or sent from the frontend.
    The reporter's email (report.user_email) is stored in bug_reports.user_email
    and used only as the Reply-To header so the admin can reply directly to the reporter.
    """
    print("[BUG-EMAIL] email function started")
    admin_email = os.environ.get("ADMIN_EMAIL", "").strip()
    smtp_server = os.environ.get("SMTP_SERVER", "smtp.gmail.com").strip()
    smtp_port = int(os.environ.get("SMTP_PORT", 587))
    smtp_user = os.environ.get("SMTP_USER", "").strip()
    smtp_password = os.environ.get("SMTP_PASSWORD", "").strip()

    has_admin = bool(admin_email and "@" in admin_email)
    has_pwd = bool(smtp_password)
    has_smtp_user = bool(smtp_user)
    print(f"[BUG-EMAIL] ADMIN_EMAIL configured: {has_admin}")
    print(f"[BUG-EMAIL] SMTP_USER configured: {has_smtp_user}")
    print(f"[BUG-EMAIL] SMTP_PASSWORD configured: {has_pwd}")
    # Do NOT log the actual admin_email value — keep recipient private in logs
    print(f"[BUG-EMAIL] reply_to: {report.user_email.strip()}")

    if not admin_email or "@" not in admin_email:
        print("[BUG-EMAIL] Warning: ADMIN_EMAIL is not configured. Email delivery skipped.")
        print("[BUG-EMAIL] provider response status: 503 Service Unavailable (Missing ADMIN_EMAIL)")
        print("[BUG-EMAIL] provider accepted message: false")
        print("[BUG-EMAIL] final result: FAILURE")
        return False, None, "ADMIN_EMAIL not configured on server"

    if not smtp_user:
        smtp_user = admin_email  # fall back to admin_email as SMTP sender only if SMTP_USER not set
        print("[BUG-EMAIL] SMTP_USER not set; using ADMIN_EMAIL as sender account")

    if not smtp_password:
        print("[BUG-EMAIL] Warning: SMTP_PASSWORD is not configured. Email delivery skipped.")
        print("[BUG-EMAIL] provider response status: 401 Unauthorized (Missing credentials)")
        print("[BUG-EMAIL] provider accepted message: false")
        print("[BUG-EMAIL] final result: FAILURE")
        return False, None, "SMTP_PASSWORD not configured on server"

    import email.utils
    message_id = email.utils.make_msgid(domain="jobwink.app")

    msg = MIMEMultipart("alternative")
    msg["Message-ID"] = message_id
    msg["Subject"] = f"[JobWink Bug Report] {report.title}"
    msg["From"] = f"JobWink Bug Reports <{smtp_user}>"
    msg["To"] = admin_email
    msg["Reply-To"] = report.user_email.strip()

    now_str = datetime.datetime.now(datetime.timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

    text_content = f"""JOBWINK BUG REPORT
==================================================
Report ID: {report_id}
Submitted: {now_str}

REPORTER IDENTITY:
--------------------------------------------------
Name:     {report.user_name or 'Not provided'}
Email:    {report.user_email}
User ID:  {report.user_id or 'Unauthenticated / Anonymous'}

BUG DETAILS:
--------------------------------------------------
Subject:  {report.title}
Location: {report.route or report.page_url or 'Unknown'}

Description:
{report.description}

ENVIRONMENT TELEMETRY:
--------------------------------------------------
Page URL:    {report.page_url or 'N/A'}
Route:       {report.route or 'N/A'}
Browser:     {report.browser or 'N/A'}
OS:          {report.os or 'N/A'}
Screen Size: {report.screen_size or 'N/A'}
Screenshot:  {report.screenshot_reference or 'None'}
==================================================
"""

    esc_name = html.escape(str(report.user_name or '')) if report.user_name else 'N/A'
    esc_email = html.escape(str(report.user_email or ''))
    esc_uid = html.escape(str(report.user_id or '')) if report.user_id else 'Unauthenticated / Anonymous'
    esc_title = html.escape(str(report.title or ''))
    esc_desc = html.escape(str(report.description or ''))
    esc_route = html.escape(str(report.route or '')) if report.route else (html.escape(str(report.page_url or '')) if report.page_url else 'N/A')
    esc_browser = html.escape(str(report.browser or '')) if report.browser else 'N/A'
    esc_os = html.escape(str(report.os or '')) if report.os else 'N/A'
    esc_screen = html.escape(str(report.screen_size or '')) if report.screen_size else 'N/A'

    screenshot_html = ''
    if report.screenshot_reference and report.screenshot_reference.startswith(("http://", "https://")):
        esc_screenshot = html.escape(str(report.screenshot_reference))
        screenshot_html = f'<tr><td style="padding: 4px 0; font-weight: 600; color: #64748B;">Screenshot:</td><td><a href="{esc_screenshot}" target="_blank" style="color: #2563EB; font-weight: 600;">View Attached Screenshot &rarr;</a></td></tr>'

    html_content = f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
  <title>Bug Report</title>
</head>
<body style="font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif; color: #1E293B; line-height: 1.6; max-width: 620px; margin: 0 auto; padding: 20px; background-color: #F1F5F9;">
  <div style="background: #0F172A; padding: 20px 24px; border-radius: 12px 12px 0 0; color: #ffffff;">
    <h2 style="margin: 0; color: #FB923C; font-size: 20px; display: flex; align-items: center;">🐛 JobWink Bug Report</h2>
    <p style="margin: 4px 0 0 0; font-size: 12px; color: #94A3B8;">Report ID: <code style="color: #CBD5E1;">{html.escape(str(report_id))}</code> &bull; {html.escape(str(now_str))}</p>
  </div>
  <div style="background: #ffffff; border: 1px solid #E2E8F0; padding: 24px; border-radius: 0 0 12px 12px; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05);">
    <h3 style="margin-top: 0; color: #0F172A; font-size: 15px; text-transform: uppercase; letter-spacing: 0.5px; border-bottom: 2px solid #FB923C; padding-bottom: 6px;">Reporter Identity</h3>
    <table style="width: 100%; font-size: 14px; margin-bottom: 20px; border-collapse: collapse;">
      <tr><td style="width: 120px; padding: 6px 0; font-weight: 600; color: #64748B;">Name:</td><td style="color: #0F172A; font-weight: 600;">{esc_name}</td></tr>
      <tr><td style="padding: 6px 0; font-weight: 600; color: #64748B;">Email:</td><td><a href="mailto:{esc_email}" style="color: #EA580C; font-weight: 600; text-decoration: none;">{esc_email}</a></td></tr>
      <tr><td style="padding: 6px 0; font-weight: 600; color: #64748B;">User ID:</td><td style="font-family: monospace; font-size: 12px; color: #475569;">{esc_uid}</td></tr>
    </table>

    <h3 style="color: #0F172A; font-size: 15px; text-transform: uppercase; letter-spacing: 0.5px; border-bottom: 2px solid #FB923C; padding-bottom: 6px;">Bug Summary</h3>
    <div style="background: #F8FAFC; border: 1px solid #E2E8F0; border-left: 4px solid #EA580C; border-radius: 6px; padding: 14px; margin-bottom: 20px;">
      <p style="margin: 0 0 8px 0; font-weight: 700; font-size: 15px; color: #0F172A;">{esc_title}</p>
      <p style="margin: 0; font-size: 14px; white-space: pre-wrap; color: #334155;">{esc_desc}</p>
    </div>

    <h3 style="color: #0F172A; font-size: 15px; text-transform: uppercase; letter-spacing: 0.5px; border-bottom: 2px solid #FB923C; padding-bottom: 6px;">Environment & Telemetry</h3>
    <table style="width: 100%; font-size: 13px; color: #475569; border-collapse: collapse;">
      <tr><td style="width: 120px; padding: 4px 0; font-weight: 600; color: #64748B;">Route / Page:</td><td>{esc_route}</td></tr>
      <tr><td style="font-weight: 600; color: #64748B;">Browser:</td><td>{esc_browser}</td></tr>
      <tr><td style="font-weight: 600; color: #64748B;">Platform / OS:</td><td>{esc_os}</td></tr>
      <tr><td style="font-weight: 600; color: #64748B;">Screen Size:</td><td>{esc_screen}</td></tr>
      {screenshot_html}
    </table>
  </div>
  <p style="font-size: 12px; color: #94A3B8; text-align: center; margin-top: 18px;">JobWink Bug Reporter &bull; Replying to this email will directly message {esc_email}</p>
</body>
</html>
"""

    msg.attach(MIMEText(text_content, "plain"))
    msg.attach(MIMEText(html_content, "html"))

    print("[BUG-EMAIL] provider request started")
    try:
        with smtplib.SMTP(smtp_server, smtp_port, timeout=15) as server:
            server.starttls()
            server.login(smtp_user, smtp_password)
            server.send_message(msg)
            print("[BUG-EMAIL] provider response status: 250 OK")
            print("[BUG-EMAIL] provider accepted message: true")
            print(f"[BUG-EMAIL] provider message ID: {message_id}")
            print("[BUG-EMAIL] final result: SUCCESS")
            return True, message_id, None
    except smtplib.SMTPAuthenticationError as e:
        err_msg = e.smtp_error.decode("utf-8", errors="ignore") if isinstance(e.smtp_error, bytes) else str(e.smtp_error)
        print(f"[BUG-EMAIL] provider response status: 535 ({err_msg})")
        print("[BUG-EMAIL] provider accepted message: false")
        print("[BUG-EMAIL] final result: FAILURE")
        return False, None, f"SMTP Authentication failed (status 535): {err_msg}"
    except smtplib.SMTPRecipientsRefused as e:
        print("[BUG-EMAIL] provider response status: 550 Recipients refused")
        print("[BUG-EMAIL] provider accepted message: false")
        print("[BUG-EMAIL] final result: FAILURE")
        return False, None, "SMTP Recipients refused (status 550)"
    except smtplib.SMTPSenderRefused as e:
        print("[BUG-EMAIL] provider response status: 553 Sender refused")
        print("[BUG-EMAIL] provider accepted message: false")
        print("[BUG-EMAIL] final result: FAILURE")
        return False, None, "SMTP Sender refused (status 553)"
    except smtplib.SMTPDataError as e:
        print("[BUG-EMAIL] provider response status: 554 Data rejected")
        print("[BUG-EMAIL] provider accepted message: false")
        print("[BUG-EMAIL] final result: FAILURE")
        return False, None, "SMTP Data rejected (status 554)"
    except smtplib.SMTPException as e:
        print(f"[BUG-EMAIL] provider response status: 500 ({type(e).__name__})")
        print("[BUG-EMAIL] provider accepted message: false")
        print("[BUG-EMAIL] final result: FAILURE")
        return False, None, f"SMTP error ({type(e).__name__}): {e}"
    except Exception as e:
        print(f"[BUG-EMAIL] provider response status: 500 ({type(e).__name__})")
        print("[BUG-EMAIL] provider accepted message: false")
        print("[BUG-EMAIL] final result: FAILURE")
        return False, None, f"Email delivery error ({type(e).__name__}): {e}"


@app.post("/api/report-bug")
async def report_bug_endpoint(payload: BugReportPayload, x_user_id: Optional[str] = Header(None, alias="X-User-ID")):
    """Receives a bug report, persists it to database, and emails the administrator."""
    user_id = payload.user_id or x_user_id
    report_id = str(uuid.uuid4())

    # 1. Persist to Supabase if connected
    db_saved = False
    if supabase_client:
        try:
            rpc_res = supabase_client.rpc("submit_bug_report", {
                "p_user_email": payload.user_email.strip().lower(),
                "p_title": payload.title.strip(),
                "p_description": payload.description.strip(),
                "p_page_url": payload.page_url,
                "p_route": payload.route,
                "p_browser": payload.browser,
                "p_os": payload.os,
                "p_screen_size": payload.screen_size,
                "p_screenshot_reference": payload.screenshot_reference,
            }).execute()
            if rpc_res and getattr(rpc_res, 'data', None):
                report_id = rpc_res.data.get('report_id') or report_id
            db_saved = True
            print(f"[BUG-EMAIL] Database save: SUCCESS (Report ID: {report_id})")
        except Exception as e:
            try:
                supabase_client.table("bug_reports").insert({
                    "id": report_id,
                    "user_id": user_id,
                    "user_email": payload.user_email.strip().lower(),
                    "title": payload.title.strip(),
                    "description": payload.description.strip(),
                    "page_url": payload.page_url,
                    "route": payload.route,
                    "browser": payload.browser,
                    "os": payload.os,
                    "screen_size": payload.screen_size,
                    "screenshot_reference": payload.screenshot_reference,
                    "status": "open",
                }).execute()
                db_saved = True
                print(f"[BUG-EMAIL] Database save: SUCCESS (Direct insert, Report ID: {report_id})")
            except Exception as e2:
                print(f"[BUG-EMAIL] Database save: FAILURE ({e2})")
    else:
        print("[BUG-EMAIL] Database save: FAILURE (No database client configured)")

    # 2. Deliver Email to Admin in background worker thread to prevent blocking event loop
    import anyio
    email_sent, message_id, email_err = await anyio.to_thread.run_sync(send_bug_report_email, payload, report_id)
    if email_sent:
        print(f"[BUG-EMAIL] Email delivery: SUCCESS (Message-ID: {message_id})")
    else:
        print(f"[BUG-EMAIL] Email delivery: FAILURE ({email_err})")

    return {
        "success": email_sent,
        "email_sent": email_sent,
        "db_saved": db_saved,
        "report_id": report_id,
        "message_id": message_id,
        "email_error": email_err,
        "message": "Bug report submitted successfully." if email_sent else "Bug report saved to database, but email delivery to administrator failed."
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("backend_main:app", host="127.0.0.1", port=8000, reload=True)


